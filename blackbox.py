"""
AI 智能批改系统 — 上传作答文档 + 选择题号，获取 AI 评分、分析和讲解视频。
"""
import json
import os
import glob
import re
import tempfile
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document as DocxDocument
from openai import OpenAI

load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50MB

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
DATA_DIR = os.path.join(BASE_DIR, "data")

# ── OpenAI client ──────────────────────────────────────────────
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_BASE_URL = os.environ.get("OPENAI_BASE_URL", "")
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

openai_client = None
if OPENAI_API_KEY:
    kwargs = {"api_key": OPENAI_API_KEY}
    if OPENAI_BASE_URL:
        kwargs["base_url"] = OPENAI_BASE_URL
    openai_client = OpenAI(**kwargs)


# ── Video mapping ──────────────────────────────────────────────
def build_video_map():
    vmap = {}
    for f in glob.glob(os.path.join(OUTPUT_DIR, "*.mp4")):
        name = os.path.basename(f)
        parts = name.rsplit("_", 1)
        if len(parts) == 2 and parts[1].endswith(".mp4"):
            qnum = parts[1].replace(".mp4", "")
            vmap[qnum] = f
    return vmap

VIDEO_MAP = build_video_map()


# ── Question data ──────────────────────────────────────────────
def load_all_questions():
    questions = {}
    titles = {}
    for fname in glob.glob(os.path.join(DATA_DIR, "Q14.*.json")):
        with open(fname, encoding="utf-8") as f:
            bank = json.load(f)
        q = list(bank.values())[0]
        qid = q.get("id", "")
        questions[qid] = q
        titles[qid] = q.get("title_cn", q.get("title", ""))
    return questions, titles

QUESTIONS, TITLES = load_all_questions()


# ── Question number normalization ──────────────────────────────
def normalize_question(qnum):
    qnum = qnum.strip().replace("_", ".")
    if "." not in qnum and qnum.isdigit():
        qnum = f"14.{qnum}"
    return qnum


def find_video(qnum):
    path = VIDEO_MAP.get(qnum)
    if path:
        return path, qnum
    for k, v in VIDEO_MAP.items():
        if k.endswith(f".{qnum}") or k == qnum:
            return v, k
    return None, qnum


# ── DOCX text extraction (with OMML math → LaTeX) ─────────────
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
}

# Unicode → LaTeX Greek/symbol mapping
_UNICODE_TO_LATEX = {
    0x03B1: r'\alpha', 0x03B2: r'\beta', 0x03B3: r'\gamma', 0x03B4: r'\delta',
    0x03B5: r'\epsilon', 0x03B6: r'\zeta', 0x03B7: r'\eta', 0x03B8: r'\theta',
    0x03B9: r'\iota', 0x03BA: r'\kappa', 0x03BB: r'\lambda', 0x03BC: r'\mu',
    0x03BD: r'\nu', 0x03BE: r'\xi', 0x03C0: r'\pi', 0x03C1: r'\rho',
    0x03C3: r'\sigma', 0x03C4: r'\tau', 0x03C5: r'\upsilon', 0x03C6: r'\phi',
    0x03C7: r'\chi', 0x03C8: r'\psi', 0x03C9: r'\omega',
    0x0393: r'\Gamma', 0x0394: r'\Delta', 0x0398: r'\Theta', 0x039B: r'\Lambda',
    0x039E: r'\Xi', 0x03A0: r'\Pi', 0x03A3: r'\Sigma', 0x03A6: r'\Phi',
    0x03A8: r'\Psi', 0x03A9: r'\Omega',
    0x221E: r'\infty', 0x2211: r'\sum', 0x220F: r'\prod', 0x222B: r'\int',
    0x2202: r'\partial', 0x2207: r'\nabla', 0x221A: r'\sqrt',
    0x00D7: r'\times', 0x00F7: r'\div', 0x00B1: r'\pm', 0x2213: r'\mp',
    0x2264: r'\leq', 0x2265: r'\geq', 0x2260: r'\neq', 0x2248: r'\approx',
    0x2192: r'\to', 0x2190: r'\leftarrow', 0x21D2: r'\Rightarrow',
    0x22C5: r'\cdot', 0x2026: r'\ldots', 0x22EF: r'\cdots',
    0x03D5: r'\phi', 0x00B0: r'\circ',  # φ (straight phi), ° (degree)
}

def _text_to_latex(text):
    """Convert Unicode characters in text to LaTeX equivalents."""
    if not text:
        return ''
    result = []
    for ch in text:
        code = ord(ch)
        if code in _UNICODE_TO_LATEX:
            result.append(_UNICODE_TO_LATEX[code])
        else:
            result.append(ch)
    return ''.join(result)


def _omml_to_latex(elem):
    """Recursively convert an OMML element to LaTeX string."""
    tag = elem.tag
    local = tag.split('}')[-1] if '}' in tag else tag

    if local == 'oMath' or local == 'oMathPara':
        parts = []
        for child in elem:
            s = _omml_to_latex(child)
            if s:
                parts.append(s)
        return ' '.join(parts)

    elif local == 'sSub':  # subscript: base_{sub}
        base = _get_child_content(elem, 'e')
        sub = _get_child_content(elem, 'sub')
        return f'{base}_{{{sub}}}'

    elif local == 'sSup':  # superscript: base^{sup}
        base = _get_child_content(elem, 'e')
        sup = _get_child_content(elem, 'sup')
        return f'{base}^{{{sup}}}'

    elif local == 'sSubSup':  # sub+sup: base_{sub}^{sup}
        base = _get_child_content(elem, 'e')
        sub = _get_child_content(elem, 'sub')
        sup = _get_child_content(elem, 'sup')
        return f'{base}_{{{sub}}}^{{{sup}}}'

    elif local == 'f':  # fraction: \frac{num}{den}
        num = _get_child_content(elem, 'num')
        den = _get_child_content(elem, 'den')
        return f'\\frac{{{num}}}{{{den}}}'

    elif local == 'rad':  # radical: \sqrt[index]{radicand}
        deg = _get_child_content(elem, 'deg')
        rad = _get_child_content(elem, 'e')
        if deg:
            return f'\\sqrt[{deg}]{{{rad}}}'
        return f'\\sqrt{{{rad}}}'

    elif local == 'd':  # delimiter (parentheses)
        content = _get_child_content(elem, 'e')
        return f'\\left({content}\\right)'

    elif local == 'nary':  # integral/sum/product
        # Get operator character
        nary_pr = elem.find(f'{{{NSMAP["m"]}}}naryPr')
        op_char = '∫'
        if nary_pr is not None:
            chr_elem = nary_pr.find(f'{{{NSMAP["m"]}}}chr')
            if chr_elem is not None:
                op_char = chr_elem.get(f'{{{NSMAP["m"]}}}val', '∫')
        lower = _get_child_content(elem, 'sub')
        upper = _get_child_content(elem, 'sup')
        body = _get_child_content(elem, 'e')
        op = _text_to_latex(op_char)
        result = op
        if lower or upper:
            result += f'_{{{lower}}}^{{{upper}}}'
        result += f' {body}'
        return result

    elif local == 'r':  # text run
        text = ''
        for t in elem.iter(f'{{{NSMAP["m"]}}}t'):
            if t.text:
                text += t.text
        # Also check w:t
        for t in elem.iter(f'{{{NSMAP["w"]}}}t'):
            if t.text:
                text += t.text
        return _text_to_latex(text)

    elif local == 't':  # direct text
        return _text_to_latex(elem.text or '')

    elif local in ('e', 'sub', 'sup', 'num', 'den'):
        return _get_child_content(elem, local)

    elif local in ('sSubPr', 'sSupPr', 'sSubSupPr', 'dPr', 'fPr',
                    'radPr', 'naryPr', 'ctrlPr', 'rPr', 'eqArrPr',
                    'mPr', 'boxPr', 'borderBoxPr', 'barPr', 'groupChrPr',
                    'limLowPr', 'limUppPr', 'accPr', 'sPre', 'funcPr',
                    'w_rPr', 'w_rFonts', 'w_i', 'w_sz', 'w_szCs',
                    'w_bookmarkStart', 'w_bookmarkEnd'):
        return ''

    else:
        # Unknown element: recurse into children
        parts = []
        for child in elem:
            s = _omml_to_latex(child)
            if s:
                parts.append(s)
        return ' '.join(parts)


def _get_child_content(elem, child_name):
    """Get concatenated LaTeX content of all children with given name."""
    ns = NSMAP['m']
    children = elem.findall(f'{{{ns}}}{child_name}')
    if not children:
        return ''
    parts = []
    for child in children:
        for grandchild in child:
            s = _omml_to_latex(grandchild)
            if s:
                parts.append(s)
    return ' '.join(parts)


def _extract_paragraph_with_math(para):
    """Extract text from a paragraph, converting OMML math to LaTeX."""
    parts = []
    for child in para._element:
        tag = child.tag
        local = tag.split('}')[-1] if '}' in tag else tag
        if tag == f'{{{NSMAP["w"]}}}r':
            for t in child.iter(f'{{{NSMAP["w"]}}}t'):
                if t.text:
                    parts.append(t.text)
        elif local == 'oMath':
            latex = _omml_to_latex(child).strip()
            if latex:
                parts.append(f' ${latex}$ ')
        elif local == 'oMathPara':
            latex = _omml_to_latex(child).strip()
            if latex:
                parts.append(f'\n$$ {latex} $$\n')
    return ''.join(parts)


def extract_docx_text(filepath):
    try:
        doc = DocxDocument(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = _extract_paragraph_with_math(para).strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = _extract_paragraph_with_math(para).strip()
                        if text:
                            paragraphs.append(text)
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[文档解析失败: {e}]"


# ── Build grading prompt ──────────────────────────────────────
def build_grading_prompt(question_data, student_text):
    q = question_data
    title = q.get("title_cn", q.get("title", ""))
    desc = q.get("description_cn", q.get("description", ""))

    # Build reference answer text
    ref_steps = []
    for step in q.get("steps", []):
        step_text = f"步骤{step['id']}: {step.get('title_cn', step.get('title', ''))}"
        if step.get("formula_cn"):
            step_text += f"\n  公式: {step['formula_cn']}"
        if step.get("explanation_cn"):
            step_text += f"\n  说明: {step['explanation_cn']}"
        if step.get("result"):
            step_text += f"\n  结果: {step.get('result_cn', step['result'])}"
        ref_steps.append(step_text)

    ref_answer = "\n".join(ref_steps)
    concepts = ", ".join(q.get("key_concepts_cn", []))

    prompt = f"""你是一位微波与智能天线课程的助教。请对比学生的作答和参考答案，给出评分和分析。

题目：{title}
题干：{desc}

参考答案：
{ref_answer}

关键概念：{concepts}

学生作答：
{student_text}

请严格以以下 JSON 格式返回（不要包含任何其他文字）：
{{
  "score": 0到100的整数,
  "step_analysis": [
    {{"step": 1, "title": "步骤标题", "status": "correct或partial或wrong或missing", "comment": "对该步骤的具体评价"}},
    ...
  ],
  "overall_comment": "对学生作答的总体评价，2-3句话",
  "suggestions": ["针对性的学习建议1", "建议2", "建议3"],
  "strengths": ["学生作答中的优点1", "优点2"]
}}

评分标准：
- 正确完成所有步骤且逻辑清晰：85-100
- 大部分步骤正确但有小错误：70-84
- 部分步骤正确但有明显错误：50-69
- 大部分步骤缺失或错误：0-49

status 含义：
- correct: 该步骤完全正确
- partial: 该步骤部分正确，有小问题
- wrong: 该步骤存在明显错误
- missing: 学生未作答该步骤"""

    return prompt


# ── Call OpenAI API ────────────────────────────────────────────
def call_openai_grading(prompt):
    if not openai_client:
        # Fallback: return a mock result when no API key is set
        return {
            "score": 0,
            "step_analysis": [],
            "overall_comment": "OpenAI API key 未配置，无法进行 AI 评分。请设置环境变量 OPENAI_API_KEY。",
            "suggestions": ["请配置 OpenAI API key 后重试"],
            "strengths": [],
            "error": "NO_API_KEY"
        }

    try:
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "你是微波与智能天线课程的 AI 助教，负责批改学生作业。请始终返回有效的 JSON 格式。"},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=4000,
        )
        content = response.choices[0].message.content.strip()
        # Extract JSON from response (handle markdown code blocks)
        json_match = re.search(r'```(?:json)?\s*(.*?)\s*```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)
        return json.loads(content)
    except json.JSONDecodeError:
        return {
            "score": 0,
            "step_analysis": [],
            "overall_comment": "AI 返回结果解析失败，请重试。",
            "suggestions": ["请检查网络连接后重试"],
            "strengths": [],
            "error": "PARSE_ERROR"
        }
    except Exception as e:
        return {
            "score": 0,
            "step_analysis": [],
            "overall_comment": f"AI 评分服务异常: {str(e)}",
            "suggestions": ["请检查 API key 和网络连接"],
            "strengths": [],
            "error": str(e)
        }


# ── Format reference answer for display ────────────────────────
def format_reference_answer(question_data):
    steps = []
    for step in question_data.get("steps", []):
        s = {
            "id": step["id"],
            "title": step.get("title_cn", step.get("title", "")),
            "formula": step.get("formula", ""),  # LaTeX for rendering
            "explanation": step.get("explanation_cn", step.get("explanation", "")),
            "result": step.get("result", ""),  # LaTeX for rendering
        }
        steps.append(s)
    return {
        "title": question_data.get("title_cn", ""),
        "description": question_data.get("description_cn", ""),
        "steps": steps,
        "concepts": question_data.get("key_concepts_cn", []),
    }


# ── Routes ────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("blackbox.html", titles=TITLES)


@app.route("/api/analyze", methods=["POST"])
def analyze():
    qnum = request.form.get("question", "").strip()
    doc = request.files.get("doc")

    if not qnum:
        return jsonify({"error": "请选择题号"}), 400

    normalized = normalize_question(qnum)
    question_data = QUESTIONS.get(normalized)
    if not question_data:
        return jsonify({"error": f"未找到题号 {qnum} 的题目数据"}), 404

    # Extract student answer text
    student_text = ""
    if doc and doc.filename:
        suffix = os.path.splitext(doc.filename)[1].lower()
        if suffix in (".docx", ".doc"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                student_text = extract_docx_text(tmp.name)
                os.unlink(tmp.name)
        else:
            return jsonify({"error": "请上传 .docx 格式的文件"}), 400
    else:
        return jsonify({"error": "请上传作答文档"}), 400

    if not student_text.strip():
        return jsonify({"error": "文档内容为空"}), 400

    # Build prompt and call AI
    prompt = build_grading_prompt(question_data, student_text)
    grading_result = call_openai_grading(prompt)

    # Find video
    video_path, resolved_qnum = find_video(normalized)
    video_url = f"/video/{resolved_qnum}" if video_path else None
    download_url = f"/download/{resolved_qnum}" if video_path else None

    # Format reference answer
    reference = format_reference_answer(question_data)

    title = TITLES.get(normalized, "")

    return jsonify({
        "status": "success",
        "question": normalized,
        "title": title,
        "score": grading_result.get("score", 0),
        "step_analysis": grading_result.get("step_analysis", []),
        "overall_comment": grading_result.get("overall_comment", ""),
        "suggestions": grading_result.get("suggestions", []),
        "strengths": grading_result.get("strengths", []),
        "reference": reference,
        "video_url": video_url,
        "download_url": download_url,
        "student_text": student_text,
    })


@app.route("/video/<qnum>")
def serve_video(qnum):
    video_path = VIDEO_MAP.get(qnum)
    if not video_path or not os.path.exists(video_path):
        return "Not found", 404
    return send_file(video_path, mimetype="video/mp4")


@app.route("/download/<qnum>")
def download_video(qnum):
    video_path = VIDEO_MAP.get(qnum)
    if not video_path or not os.path.exists(video_path):
        return "Not found", 404
    title = TITLES.get(qnum, qnum)
    return send_file(video_path, mimetype="video/mp4",
                     as_attachment=True, download_name=f"{title}_{qnum}.mp4")


if __name__ == "__main__":
    print(f"Loaded {len(VIDEO_MAP)} videos, {len(QUESTIONS)} questions")
    if not OPENAI_API_KEY:
        print("WARNING: OPENAI_API_KEY not set in .env. AI grading will return mock results.")
        print("Please edit .env file and fill in your API key.")
    else:
        print(f"AI model: {OPENAI_MODEL} @ {OPENAI_BASE_URL or 'default'}")
    app.run(debug=True, port=5001)
